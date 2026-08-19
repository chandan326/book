import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx

def generate_pdf_book(book_data: dict) -> bytes:
    """
    Generates a publication-ready PDF document complete with Title Page,
    Table of Contents, Front Matter, Chapter Headings, Section Content, and Back Matter.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=28,
        leading=34,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=16,
        leading=22,
        textColor=colors.HexColor('#2563EB'),
        spaceAfter=30
    )
    
    chapter_title_style = ParagraphStyle(
        'ChapterHeader',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=20,
        spaceAfter=12
    )

    section_title_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=14,
        spaceAfter=8
    )

    body_style = ParagraphStyle(
        'BodyContent',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        leading=16,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=10
    )

    story = []

    # 1. Title Page
    story.append(Spacer(1, 100))
    story.append(Paragraph(book_data.get("title", "Untitled Book"), title_style))
    if book_data.get("subtitle"):
        story.append(Paragraph(book_data.get("subtitle"), subtitle_style))
    story.append(Spacer(1, 40))
    story.append(Paragraph(f"Author: {book_data.get('author_name', 'Author')}", body_style))
    story.append(Paragraph(f"Genre: {book_data.get('genre', 'General')}", body_style))
    story.append(PageBreak())

    # 2. Table of Contents
    story.append(Paragraph("Table of Contents", chapter_title_style))
    story.append(Spacer(1, 10))
    for chapter in book_data.get("chapters", []):
        story.append(Paragraph(f"• {chapter.get('title')}", body_style))
    story.append(PageBreak())

    # 3. Chapters & Sections
    for chapter in book_data.get("chapters", []):
        story.append(Paragraph(chapter.get("title", "Chapter"), chapter_title_style))
        if chapter.get("summary"):
            story.append(Paragraph(f"<i>Summary: {chapter.get('summary')}</i>", body_style))
        story.append(Spacer(1, 10))
        
        for section in chapter.get("sections", []):
            story.append(Paragraph(section.get("title", "Section"), section_title_style))
            content = section.get("content", "")
            for p in content.split("\n\n"):
                if p.strip():
                    story.append(Paragraph(p.strip(), body_style))
            story.append(Spacer(1, 10))
            
        story.append(PageBreak())

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx_book(book_data: dict) -> bytes:
    """
    Generates a editable DOCX manuscript file preserving heading hierarchy and paragraph formatting.
    """
    doc = docx.Document()
    
    # Title & Subtitle
    title_p = doc.add_heading(book_data.get("title", "Untitled Book"), level=0)
    if book_data.get("subtitle"):
        sub_p = doc.add_paragraph(book_data.get("subtitle"))
        sub_p.runs[0].italic = True

    doc.add_paragraph(f"Author: {book_data.get('author_name', 'Author')}")
    doc.add_page_break()

    # Table of Contents Header
    doc.add_heading("Table of Contents", level=1)
    for ch in book_data.get("chapters", []):
        doc.add_paragraph(f"• {ch.get('title')}")
    doc.add_page_break()

    # Chapters & Sections
    for ch in book_data.get("chapters", []):
        doc.add_heading(ch.get("title", "Chapter"), level=1)
        if ch.get("summary"):
            p = doc.add_paragraph(f"Summary: {ch.get('summary')}")
            p.runs[0].italic = True

        for sec in ch.get("sections", []):
            doc.add_heading(sec.get("title", "Section"), level=2)
            content = sec.get("content", "")
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
